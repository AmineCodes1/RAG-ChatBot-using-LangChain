"""Document loading utilities for PDF, TXT, and DOCX ingestion."""

from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document
from loguru import logger

from config import SUPPORTED_EXTENSIONS


def load_document(file_path: str | Path) -> list[Document]:
    """Load a file into LangChain Documents based on file extension.

    Args:
        file_path: Path to a single file to ingest.

    Returns:
        A list of LangChain Document objects generated from the file.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file extension is unsupported.
    """
    resolved_path = Path(file_path)
    if not resolved_path.exists() or not resolved_path.is_file():
        raise FileNotFoundError(f"File not found: {resolved_path}")

    extension = resolved_path.suffix.lower()
    logger.info("Loading file='{}' with extension='{}'", resolved_path, extension)

    if extension == ".pdf":
        documents = _load_pdf(resolved_path)
    elif extension == ".txt":
        documents = _load_txt(resolved_path)
    elif extension == ".docx":
        documents = _load_docx(resolved_path)
    else:
        raise ValueError(
            f"Unsupported extension '{extension}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    logger.info(
        "Loaded file='{}' extension='{}' documents={}",
        resolved_path,
        extension,
        len(documents),
    )
    return documents


def load_documents_from_dir(dir_path: str | Path) -> list[Document]:
    """Recursively load all supported files from a directory tree.

    Args:
        dir_path: Root directory containing files to ingest.

    Returns:
        A flat list of all loaded Documents across supported files.
    """
    root_dir = Path(dir_path)
    if not root_dir.exists() or not root_dir.is_dir():
        raise FileNotFoundError(f"Directory not found: {root_dir}")

    supported_extensions = {ext.lower() for ext in SUPPORTED_EXTENSIONS}
    all_documents: list[Document] = []
    failed_files: list[str] = []

    candidate_files = sorted(
        file for file in root_dir.rglob("*") if file.is_file() and file.suffix.lower() in supported_extensions
    )
    logger.info(
        "Scanning directory='{}' supported_files={}",
        root_dir,
        len(candidate_files),
    )

    for file_path in candidate_files:
        try:
            all_documents.extend(load_document(file_path))
        except Exception as exc:
            failed_files.append(str(file_path))
            logger.error("Failed loading file='{}': {}", file_path, exc)

    logger.info(
        "Directory load complete root='{}' total_documents={} failed_files={}",
        root_dir,
        len(all_documents),
        len(failed_files),
    )
    if failed_files:
        logger.warning("Failed files: {}", failed_files)

    return all_documents


def _load_pdf(file_path: Path) -> list[Document]:
    """Load a PDF into one Document per page.

    Args:
        file_path: Path to a PDF file.

    Returns:
        A list of page-level Documents.
    """
    documents: list[Document] = []
    with fitz.open(file_path) as pdf_document:
        for page_number, page in enumerate(pdf_document, start=1):
            documents.append(
                Document(
                    page_content=page.get_text("text"),
                    metadata={
                        "source": file_path.name,
                        "page": page_number,
                        "file_type": "pdf",
                    },
                )
            )
    return documents


def _load_txt(file_path: Path) -> list[Document]:
    """Load a UTF-8 text file into a single Document.

    Args:
        file_path: Path to a text file.

    Returns:
        A single-item list containing one Document for the file.
    """
    text_content = file_path.read_text(encoding="utf-8")
    return [
        Document(
            page_content=text_content,
            metadata={"source": file_path.name, "page": 0, "file_type": "txt"},
        )
    ]


def _load_docx(file_path: Path) -> list[Document]:
    """Load a DOCX file into one Document per non-empty paragraph.

    Args:
        file_path: Path to a DOCX file.

    Returns:
        A list of paragraph-level Documents.
    """
    docx_document = DocxDocument(file_path)
    documents: list[Document] = []
    for paragraph in docx_document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={"source": file_path.name, "page": 0, "file_type": "docx"},
            )
        )
    return documents
